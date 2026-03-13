import streamlit as st
import os
import io
import PyPDF2
import docx
import json
from docxtpl import DocxTemplate
from dotenv import load_dotenv
from groq import Groq

# Load environment variables
load_dotenv()

# --- Page Configuration ---
st.set_page_config(
    page_title="Solution Document Synthesizer",
    page_icon="📄",
    layout="wide" 
)

# --- Session State Initialization ---
if "generated_sd" not in st.session_state:
    st.session_state.generated_sd = ""

if "previous_sd" not in st.session_state: # Included the Undo state you asked for earlier!
    st.session_state.previous_sd = None

if "knowledge_base" not in st.session_state:
    if os.path.exists("knowledge_base.json"):
        with open("knowledge_base.json", "r", encoding="utf-8") as f:
            st.session_state.knowledge_base = json.load(f)
    else:
        st.session_state.knowledge_base = {} 

if "awaiting_missing_info" not in st.session_state:
    st.session_state.awaiting_missing_info = False

if "missing_info_report" not in st.session_state:
    st.session_state.missing_info_report = ""

if "cached_docs" not in st.session_state:
    st.session_state.cached_docs = {}

if "cr_conflict" not in st.session_state:
    st.session_state.cr_conflict = None

if "kb_edit_mode" not in st.session_state:
    st.session_state.kb_edit_mode = {}

# --- State Management Helpers ---
def clear_inputs():
    """Clears uploaded files, text inputs, and resets generation state."""
    keys_to_clear = ['reg_file', 'brd_file', 'add_file', 'current_cr', 'parent_cr']
    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]
            
    st.session_state.generated_sd = ""
    st.session_state.previous_sd = None
    st.session_state.awaiting_missing_info = False
    st.session_state.missing_info_report = ""
    st.session_state.cached_docs = {}
    st.session_state.cr_conflict = None

# --- Groq & LLM Helper Functions ---
def check_document_relevance(content, doc_type, api_key):
    client = Groq(api_key=api_key)
    validation_prompt = f"""You are an expert document classifier. 
    Task: Determine if the provided text is a relevant '{doc_type}' document.
    - If it is relevant or contains elements of a {doc_type}, respond EXACTLY with the word: VALID
    - If it is completely irrelevant (e.g., a recipe, random code, a personal letter, unrelated topic), respond with: INVALID: [Brief 1-sentence reason]
    """
    try:
        response = client.chat.completions.create(
            messages=[
                {"role": "system", "content": validation_prompt},
                {"role": "user", "content": f"Document Text:\n{content[:4000]}"} 
            ],
            model="llama-3.1-8b-instant",
            temperature=0.0,
            max_tokens=60,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"INVALID: API Error during validation - {str(e)}"

def fill_word_template(json_data_string, template_path="template.docx"):
    try:
        clean_json_string = json_data_string.strip()
        if clean_json_string.startswith("```json"):
            clean_json_string = clean_json_string[7:]
        elif clean_json_string.startswith("```"):
            clean_json_string = clean_json_string[3:]
            
        if clean_json_string.endswith("```"):
            clean_json_string = clean_json_string[:-3]
            
        context = json.loads(clean_json_string.strip(), strict = False)
        doc = DocxTemplate(template_path)
        doc.render(context)
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()
        
    except json.JSONDecodeError as e:
        st.error(f"Error: The AI did not return valid JSON. Details: {e}")
        return None
    except Exception as e:
        st.error(f"Error generating Word document: {e}")
        return None

def get_next_cr_number(kb):
    if not kb:
        return "CR000001"
    max_num = 0
    for key in kb.keys():
        if key.startswith("CR") and key[2:].isdigit():
            num = int(key[2:])
            if num > max_num:
                max_num = num
    next_num = max_num + 1
    return f"CR{next_num:06d}"

def check_missing_information(reg_text, brd_text, api_key):
    client = Groq(api_key=api_key)
    prompt = f"""You are a precise business analyst.
    Task: Review the input documents and identify if any core architectural details are completely missing.
    Check for: Business Logic, Process Flows, Technical/System Impacts, User Roles, and Data Migration.
    
    - If sufficient information is present to build a Solution Document, respond EXACTLY with the word: NONE
    - If critical information is missing, provide a concise bulleted list of the missing details.
    
    Input Documents:
    [Regulatory]: {reg_text[:3000]}
    [BRD/URF]: {brd_text[:3000]}
    """
    try:
        response = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.1-8b-instant",
            temperature=0.1,
            max_tokens=300,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return "NONE" 

def get_groq_response(system_content, user_content, api_key):
    if not api_key:
        st.error("Please set your Groq API Key in the .env file.")
        return None
    client = Groq(api_key=api_key)
    try:
        full_history = [
            {"role": "system", "content": system_content},
            {"role": "user", "content": user_content} 
        ]
        chat_completion = client.chat.completions.create(
            messages=full_history,
            model="llama-3.3-70b-versatile",
            temperature=0.4, 
            max_tokens=4096, 
        )
        return chat_completion.choices[0].message.content
    except Exception as e:
        st.error(f"API Error: {str(e)}")
        return None

def extract_text_from_files(uploaded_files):
    if not uploaded_files:
        return ""
    if not isinstance(uploaded_files, list):
        uploaded_files = [uploaded_files]
        
    combined_text = ""
    for uploaded_file in uploaded_files:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        combined_text += f"\n\n--- Content from {uploaded_file.name} ---\n\n"
        
        try:
            if file_extension == 'txt':
                stringio = io.StringIO(uploaded_file.getvalue().decode("utf-8"))
                combined_text += stringio.read()
            elif file_extension == 'pdf':
                reader = PyPDF2.PdfReader(uploaded_file)
                text_blocks = [page.extract_text() for page in reader.pages if page.extract_text()]
                combined_text += "\n".join(text_blocks)
            elif file_extension == 'docx':
                doc = docx.Document(uploaded_file)
                combined_text += "\n".join([para.text for para in doc.paragraphs])
            else:
                st.error(f"Unsupported file type: {file_extension}")
        except Exception as e:
            st.error(f"Error reading {uploaded_file.name}: {e}")
            
    return combined_text.strip()

def refine_solution_document(current_sd, edit_instruction, api_key):
    if not api_key:
        st.error("Please set your Groq API Key.")
        return None
        
    client = Groq(api_key=api_key)
    
    system_prompt = """You are an expert AI Solution Document Architect. 
    Your task is to revise the provided JSON Solution Document based strictly on the user's instructions. 
    CRITICAL: You MUST return the output as a valid, raw JSON object exactly matching the keys of the original document. Do not wrap it in markdown. Do not include introductory text."""
    
    user_prompt = f"### CURRENT JSON DOCUMENT:\n{current_sd}\n\n### REVISION INSTRUCTIONS:\n{edit_instruction}"
    
    try:
        response = client.chat.completions.create(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt} 
            ],
            model="llama-3.3-70b-versatile",
            temperature=0.3,
            max_tokens=4096,
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"API Error during revision: {str(e)}")
        return None

def execute_synthesis_pipeline(reg, brd, additional, api_key):
    synthesize_prompt = f"""
I am providing several key documents below to be synthesized into a single Solution Document (SD) JSON object.

Please synthesize them exactly according to the provided system instructions.

---
### INPUT DOCUMENTS:
**1. Regulatory Document Content:**
{reg if reg.strip() else "None provided. Proceed with synthesis without mapping regulatory constraints."}

**2. Business Requirement Document (BRD/URF) Content:**
{brd}
"""
    if additional.strip():
        synthesize_prompt += f"\n---\n**3. Additional Supporting Document Content:**\n{additional}\n"
        
    if parent_sd_content:
            synthesize_prompt += f"""
---
**4. Parent Solution Document (Reference):**
{parent_sd_content}
*Instruction: Use this older Parent SD to pre-fill or carry over common project details.*
"""

    with st.status("🛠️ Synthesizing Final Solution Document...", expanded=True) as status:
        st.write("Mapping Regulatory constraints to Business requirements...")
        st.write("Extracting JSON properties...")
        
        response = get_groq_response(SYSTEM_PROMPT, synthesize_prompt, api_key)
        
        st.write("Finalizing synthesized JSON...")
        status.update(label="Synthesis Complete!", state="complete", expanded=False)
        
        if response:
            st.session_state.generated_sd = response
            st.session_state.awaiting_missing_info = False
            st.rerun()

# --- Constants & System Prompt ---
SYSTEM_PROMPT = """
You are an expert AI Solution Document Architect specializing STRICTLY in Core Banking Systems (CBS).

YOUR MANDATE AND STRICT RULES:
1.  **ABSOLUTE SCOPE RESTRICTION (CBS ONLY - CRITICAL):** You must ONLY extract, synthesize, and include development requirements, process flows, or changes that occur explicitly WITHIN the CBS. Ignore all external portals, UI/UX, or non-CBS systems.
2.  **STRICT JSON OUTPUT REQUIRED:** You MUST output your response as a valid, raw JSON object. DO NOT wrap the JSON in markdown formatting. DO NOT include introductory text. MUST strictly escape all inner quotes (\\") and avoid literal unescaped line breaks inside strings.
3.  **JSON SCHEMA (MANDATORY KEYS):** Your JSON keys must exactly match the list below. Do not add or remove any keys. If information is missing for a key, output "Information not provided in source documents."
    {
        "cr_number": "Extract from input or leave blank for auto-generation",
        "month_year": "Current month and year",
        "module_name": "Name of the relevant CBS module",
        "functionality_name": "Name of the specific functionality being changed",
        "brief_description": "A 1-2 sentence high-level summary",
        "cr_details": "Detailed overview of the Change Request",
        "scope_of_change": "In-depth explanation of the scope of CBS changes",
        "executive_summary": "Comprehensive executive summary mapping business needs to CBS constraints",
        "existing_functionality": "Detailed explanation of the current process flow",
        "technical_feasibility": "Analysis of technical feasibility",
        "proposed_solution_details": "Deep, exhaustive explanation of the CBS solution, business logic, and architectural changes.",
        "assumptions": "List of assumptions",
        "limitations": "List of technical or business limitations",
        "user_type_specifications": "Required user types and capabilities",
        "maker_checker_specifications": "Maker and Checker rules",
        "data_migration": "Data migration applicability and details",
        "implementation_plan": "Step-by-step implementation plan",
        "archival_policy": "Archival policy and performance testing requirements",
        "business_acceptance_scenario": "Business acceptance and UAT scenarios",
        "references": "Any references or 'None provided'",
        "abbr_1_term": "1st abbreviation found in text (or 'N/A' if none)",
        "abbr_1_def": "Definition of 1st abbreviation (or 'N/A')",
        "abbr_2_term": "2nd abbreviation found in text (or 'N/A')",
        "abbr_2_def": "Definition of 2nd abbreviation (or 'N/A')",
        "abbr_3_term": "3rd abbreviation found in text (or 'N/A')",
        "abbr_3_def": "Definition of 3rd abbreviation (or 'N/A')"
    }
4.  **Depth and Comprehensiveness (FOR CBS ONLY):** Within the JSON values, elaborate extensively. Provide highly detailed explanations of the CBS business logic and technical implementations. If you need line breaks inside your text, you MUST use the literal escaped string "\\n" (backslash n). Do not use actual unescaped line breaks or carriage returns.
5.  **STRICT GROUNDING:** ONLY use the information explicitly provided in the uploaded documents. Do not invent outside the context.
6. **Detailed Solutioning:** For the "proposed_solution_details" key, provide an extremely deep and comprehensive explanation of the proposed CBS solution. Include detailed descriptions of the business logic, technical architecture, and any changes to existing processes. This section should be thorough enough to guide a development team in understanding the full scope of the solution. It should be 1000 words or more if the information is available in the source documents. If certain details are missing, clearly state that they were not provided in the source materials. The goal is to create a rich, detailed narrative that captures all aspects of the proposed solution within the CBS context.
7. **Executive Summary Depth:** For the "executive_summary" key, provide a comprehensive summary that not only maps the business needs to CBS constraints but also highlights the key architectural decisions and their rationale. This should give executives a clear understanding of the strategic implications of the change.
"""
# --- Sidebar: Configuration ---
with st.sidebar:
    st.header("⚙️ Configuration")
    api_key = os.getenv("GROQ_API_KEY")
    
    if api_key:
        st.success("API Key loaded", icon="✅")
    else:
        st.error("No API Key found. Please set GROQ_API_KEY in your .env file.")
    
    st.markdown("---")
    st.markdown("### 📝 Instructions")
    st.markdown("1. Ensure `template.docx` is present in the app directory.")
    st.markdown("2. Upload all mandatory input documents below.")
    st.markdown("3. Click 'Analyze Documents'.")
    st.markdown("4. Supply missing info (if prompted) and Generate.")
    
    if st.button("Clear Inputs & Reset", key="clear_all_btn"):
        clear_inputs()
        st.rerun()

    st.markdown("---")
    st.markdown("### 🗄️ Legacy Import")
    with st.container(border=True):
        st.caption("Insert Old SD into Knowledge base")
        st.file_uploader("Upload existing SD files (TBD Later)", disabled=True)

# --- Main Interface Design ---
st.title("📄 Solution Document Synthesizer")
st.subheader("Architecture-Driven Document Pipeline")
st.divider()

# --- 1. Input Section (Always Visible) ---
st.header("Input Section") 
st.info("ℹ️ Using dynamic JSON structure to inject into `template.docx`.")

col_cr1, col_cr2 = st.columns(2)
with col_cr1:
    auto_cr = get_next_cr_number(st.session_state.knowledge_base)
    current_cr = st.text_input("Current CR No. (Auto-generated)", value=auto_cr, key="current_cr")
with col_cr2:
    parent_cr = st.text_input("Parent CR No. (Optional)", key="parent_cr")

parent_sd_content = ""
if parent_cr and parent_cr in st.session_state.knowledge_base:
    st.success(f"✅ Parent CR '{parent_cr}' found! Details will be pre-filled as a reference.")
    parent_sd_content = st.session_state.knowledge_base[parent_cr]
elif parent_cr:
    st.warning(f"⚠️ Parent CR '{parent_cr}' not found in Knowledge Base.")

colA, colB = st.columns(2)
with colA:
    reg_files = st.file_uploader("Upload Regulatory Document(s) (Optional)", type=['txt', 'pdf', 'docx'], key="reg_file", accept_multiple_files=True)
with colB:
    brd_files = st.file_uploader("Upload BRD / URF Document(s)", type=['txt', 'pdf', 'docx'], key="brd_file", accept_multiple_files=True)

st.markdown("---")
add_files = st.file_uploader("Upload Additional Document(s) (Supporting - Optional)", type=['txt', 'pdf', 'docx'], key="add_file", accept_multiple_files=True)
st.markdown("<br>", unsafe_allow_html=True)

if not st.session_state.awaiting_missing_info:
    generate_btn = st.button("🚀 Analyze Documents & Generate SD", type="primary", use_container_width=True)

    if generate_btn:
        reg_content = extract_text_from_files(reg_files)
        brd_content = extract_text_from_files(brd_files)
        additional_final_content = extract_text_from_files(add_files)

        if not brd_content.strip():
            st.warning("Please upload BRD/URF input documents.")
        elif not api_key:
            st.warning("GROQ_API_KEY is not configured in your .env file.")
        else:
            with st.status("🔍 Analyzing input documents...", expanded=True) as val_status:
                if reg_content.strip():
                    st.write("Validating Regulatory Document relevance...")
                    reg_validation = check_document_relevance(reg_content, "Regulatory or Compliance", api_key)
                    if reg_validation.startswith("INVALID"):
                        val_status.update(label="Validation Failed", state="error", expanded=True)
                        st.error(f"**Regulatory Document Error:** {reg_validation.replace('INVALID:', '').strip()}")
                        st.stop()
                else:
                    st.write("No Regulatory Document provided. Skipping regulatory validation...")
                
                st.write("Validating BRD / URF Document relevance...")
                brd_validation = check_document_relevance(brd_content, "Business Requirement or Use Case", api_key)
                if brd_validation.startswith("INVALID"):
                    val_status.update(label="Validation Failed", state="error", expanded=True)
                    st.error(f"**BRD/URF Document Error:** {brd_validation.replace('INVALID:', '').strip()}")
                    st.stop()
                    
                st.write("Cross-referencing inputs with required details...")
                missing_report = check_missing_information(reg_content, brd_content, api_key)
                val_status.update(label="Analysis Complete!", state="complete", expanded=False)

            if missing_report != "NONE":
                st.session_state.awaiting_missing_info = True
                st.session_state.missing_info_report = missing_report
                st.session_state.cached_docs = {
                    "reg": reg_content,
                    "brd": brd_content,
                    "add": additional_final_content
                }
                st.rerun()
            else:
                execute_synthesis_pipeline(reg_content, brd_content, additional_final_content, api_key)

def display_human_readable_doc(json_data_string):
    """Converts the JSON data into a clean, human-readable document preview."""
    try:
        # Clean the string
        clean_json_string = json_data_string.strip()
        if clean_json_string.startswith("```json"):
            clean_json_string = clean_json_string[7:]
        elif clean_json_string.startswith("```"):
            clean_json_string = clean_json_string[3:]
        if clean_json_string.endswith("```"):
            clean_json_string = clean_json_string[:-3]
            
        data = json.loads(clean_json_string, strict=False)
        
        # Build the Header Preview
        st.markdown(f"### 📄 CR No: {data.get('cr_number', 'N/A')} | {data.get('month_year', 'N/A')}")
        st.markdown(f"**Module:** {data.get('module_name', 'N/A')} | **Functionality:** {data.get('functionality_name', 'N/A')}")
        st.info(f"**Description:** {data.get('brief_description', 'N/A')}")
        st.markdown("---")
        
        # Map the JSON keys to friendly section titles
        sections = {
            "cr_details": "1. CR Details",
            "scope_of_change": "1.1 Scope of Change",
            "executive_summary": "1.2 Executive Summary",
            "existing_functionality": "2.1 Existing Functionality with Process Flow",
            "technical_feasibility": "2.2 Technical Feasibility",
            "proposed_solution_details": "2.3 Proposed Solution Details",
            "assumptions": "2.4 Assumptions",
            "limitations": "2.5 Limitations",
            "user_type_specifications": "2.6 User Type or Capability Specifications",
            "maker_checker_specifications": "2.7 Maker Checker Specifications",
            "data_migration": "2.10 Data Migration Applicability",
            "implementation_plan": "2.11 Implementation Plan",
            "archival_policy": "2.12 Archival Policy",
            "business_acceptance_scenario": "2.13 Business Acceptance Scenario",
            "references": "2.14 References"
        }
        
        # Loop through and print only the sections that have data
        for key, title in sections.items():
            content = data.get(key, "")
            # Skip empty sections or the default "not provided" text to keep the view clean
            if content and content != "Information not provided in source documents.":
                st.markdown(f"#### {title}")
                st.markdown(content)
                
    except Exception as e:
        st.error("Could not generate document preview. Displaying raw data instead.")
        st.code(json_data_string, language="json")

# --- 1.5 Missing Information Prompt ---
if st.session_state.awaiting_missing_info and not st.session_state.generated_sd:
    st.divider()
    st.warning("⚠️ Missing Information Detected in Source Documents")
    st.markdown("We found that the following details are missing from the uploads:")
    
    # NEW: Wrap the missing info in an expander so it is collapsed by default
    with st.expander("🔍 View missing details"):
        st.info(st.session_state.missing_info_report)
    
    user_supplemental = st.text_area(
        "📝 Provide the missing information here (or leave blank to proceed without it):", 
        height=150
    )
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Proceed and Generate SD", type="primary", use_container_width=True):
            combined_additional = st.session_state.cached_docs['add']
            if user_supplemental.strip():
                combined_additional += f"\n\n[User Provided Supplemental Information]:\n{user_supplemental}"
                
            execute_synthesis_pipeline(
                st.session_state.cached_docs['reg'],
                st.session_state.cached_docs['brd'],
                combined_additional,
                api_key
            )
    with col2:
        if st.button("Cancel Generation", use_container_width=True):
            clear_inputs()
            st.rerun()

# --- 2. Post-Processing View (Output Section) ---
if st.session_state.generated_sd:
    st.divider()
    st.header("Output Section") 
    st.markdown("### Synthesized Solution Data (JSON)") 
    
    with st.container(border=True):
        # Call our new function to render the text like a document!
        display_human_readable_doc(st.session_state.generated_sd)
        
    st.markdown("### ✍️ Refine Generated Document")
    
    # Optional Undo button
    if st.session_state.previous_sd:
        if st.button("↩️ Undo Last Revision"):
            st.session_state.generated_sd = st.session_state.previous_sd
            st.session_state.previous_sd = None 
            st.success("Reverted to the previous version!")
            st.rerun()

    with st.form("edit_sd_form"):
        col_edit1, col_edit2 = st.columns([4, 1])
        with col_edit1:
            edit_instruction = st.text_input(
                "Ask the AI to change something", 
                placeholder="e.g., 'Expand the Assumptions section', 'Make the tone more formal'",
                label_visibility="collapsed"
            )
        with col_edit2:
            submit_edit = st.form_submit_button("✨ Apply Revision", use_container_width=True)
            
        if submit_edit and edit_instruction.strip():
            with st.status("🔄 Refining document...", expanded=True) as edit_status:
                st.write(f"Applying instruction: '{edit_instruction}'...")
                revised_sd = refine_solution_document(st.session_state.generated_sd, edit_instruction, api_key)
                edit_status.update(label="Revision Complete!", state="complete", expanded=False)
                
                if revised_sd:
                    st.session_state.previous_sd = st.session_state.generated_sd
                    st.session_state.generated_sd = revised_sd
                    st.rerun()
        
    st.markdown("### Post-Synthesis Pipeline")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("✅ Insert into Knowledge Base", key="save_kb_btn", type="primary", use_container_width=True):
            cr_key = st.session_state.current_cr.strip() if st.session_state.current_cr.strip() else get_next_cr_number(st.session_state.knowledge_base)
            
            if cr_key in st.session_state.knowledge_base:
                st.session_state.cr_conflict = cr_key
            else:
                st.session_state.knowledge_base[cr_key] = st.session_state.generated_sd
                with open("knowledge_base.json", "w", encoding="utf-8") as f:
                    json.dump(st.session_state.knowledge_base, f, indent=4)
                    
                st.success(f"Document added to KB under '{cr_key}'!")
                clear_inputs() 
                st.rerun()
                
        if st.session_state.cr_conflict:
            st.warning(f"⚠️ The CR No. '{st.session_state.cr_conflict}' already exists!")
            st.markdown("Would you like to auto-generate a new one or type a different one above?")
            
            res_col1, res_col2 = st.columns(2)
            with res_col1:
                if st.button("Auto-Assign & Save", use_container_width=True):
                    new_cr = get_next_cr_number(st.session_state.knowledge_base)
                    st.session_state.knowledge_base[new_cr] = st.session_state.generated_sd
                    with open("knowledge_base.json", "w", encoding="utf-8") as f:
                        json.dump(st.session_state.knowledge_base, f, indent=4)
                    
                    st.success(f"Saved successfully as '{new_cr}'!")
                    clear_inputs()
                    st.rerun()
            with res_col2:
                if st.button("I'll type a new one", use_container_width=True):
                    st.session_state.cr_conflict = None
                    st.rerun()
            
    with col2:
        word_bytes = fill_word_template(st.session_state.generated_sd)
        if word_bytes:
            st.download_button(
                label="📥 Download as Word Doc",
                data=word_bytes,
                file_name=f"{st.session_state.current_cr}_Solution_Document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
    with col3:
        if st.button("❌ Discard Output & Start Over", key="discard_btn", use_container_width=True):
            clear_inputs()
            st.rerun()

# --- 3. Knowledge Base Section ---
st.divider()
st.header("Knowledge Base")
if not st.session_state.knowledge_base:
    st.info("The knowledge base is currently empty. Generated documents can be stored here for future reference.")
else:
    for cr_key, sd_item in reversed(list(st.session_state.knowledge_base.items())):
        with st.expander(f"🗃️ CR No: {cr_key}"):
            
            col1, col2, col3, col4 = st.columns([2, 2, 2, 4])
            
            with col1:
                word_bytes = fill_word_template(sd_item)
                if word_bytes:
                    st.download_button(
                        label="📥 Word Doc",
                        data=word_bytes,
                        file_name=f"{cr_key}_Solution_Document.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_kb_{cr_key}",
                        use_container_width=True
                    )
            
            with col2:
                is_editing = st.session_state.kb_edit_mode.get(cr_key, False)
                if st.button("✏️ Edit JSON" if not is_editing else "❌ Cancel", key=f"edit_toggle_{cr_key}", use_container_width=True):
                    st.session_state.kb_edit_mode[cr_key] = not is_editing
                    st.rerun()
            
            with col3:
                if st.button("🗑️ Delete", key=f"del_{cr_key}", type="primary", use_container_width=True):
                    del st.session_state.knowledge_base[cr_key]
                    with open("knowledge_base.json", "w", encoding="utf-8") as f:
                        json.dump(st.session_state.knowledge_base, f, indent=4)
                    
                    if cr_key in st.session_state.kb_edit_mode:
                        del st.session_state.kb_edit_mode[cr_key]
                        
                    st.rerun()

            st.markdown("---")
            
            if st.session_state.kb_edit_mode.get(cr_key, False):
                new_sd_content = st.text_area("Edit JSON Document:", value=sd_item, height=400, key=f"text_area_{cr_key}")
                
                if st.button("💾 Save Changes", key=f"save_edit_{cr_key}", type="primary"):
                    st.session_state.knowledge_base[cr_key] = new_sd_content
                    with open("knowledge_base.json", "w", encoding="utf-8") as f:
                        json.dump(st.session_state.knowledge_base, f, indent=4)
                    st.session_state.kb_edit_mode[cr_key] = False
                    st.success("Changes saved!")
                    st.rerun()
            else:
                display_human_readable_doc(sd_item)
